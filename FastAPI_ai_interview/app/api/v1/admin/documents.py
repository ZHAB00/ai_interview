"""Admin document management endpoints (non-structured knowledge base)."""

import logging
import uuid
from pathlib import Path

from fastapi import APIRouter, Depends, UploadFile, File, Form
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_admin, get_db
from app.core.config import settings
from app.core.exceptions import FileTooLargeException, NotFoundException, ValidationErrorException
from app.models.document import Document
from app.models.user import User
from app.schemas.admin import DocumentListResponse, DocumentListItem, DocumentStatusResponse, DocumentUploadResponse

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/admin/documents", tags=["管理后台-文档管理"])

# Allowed extensions for knowledge base documents
_DOC_MAGIC = {
    ".pdf":  b"%PDF",
    ".docx": b"PK\x03\x04",
}
_DOC_TEXT_EXT = {".txt", ".md", ".py", ".js", ".ts", ".json", ".yaml", ".yml"}


def _check_doc_magic(filename: str, data: bytes) -> None:
    """Validate file content against its claimed extension."""
    ext = Path(filename).suffix.lower()
    if ext in _DOC_MAGIC:
        expected = _DOC_MAGIC[ext]
        if data[:len(expected)] != expected:
            raise ValidationErrorException(f"文件内容与扩展名 {ext} 不匹配")
        return
    if ext in _DOC_TEXT_EXT:
        # Ensure it's valid text, not binary garbage
        try:
            data.decode("utf-8")
        except UnicodeDecodeError:
            raise ValidationErrorException(f"文件内容不是有效的 UTF-8 文本")
        return
    raise ValidationErrorException(f"不支持的文档类型: {ext}")


@router.post("", response_model=DocumentUploadResponse, status_code=202)
async def upload_document(
    file: UploadFile = File(...),
    description: str = Form(default=""),
    tags: str = Form(default=""),
    current_user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """上传非结构化文档，用于补充知识库。上传后异步处理。"""
    if not file.filename:
        raise ValidationErrorException("文件名为空")

    file_bytes = await file.read()
    if len(file_bytes) > settings.MAX_UPLOAD_SIZE:
        raise FileTooLargeException("文件大小超过限制")

    _check_doc_magic(file.filename, file_bytes)

    # Save file
    upload_dir = Path(settings.UPLOAD_DIR) / "documents"
    upload_dir.mkdir(parents=True, exist_ok=True)
    ext = Path(file.filename).suffix.lower()
    saved_name = f"{uuid.uuid4().hex}{ext}"
    saved_path = str(upload_dir / saved_name)
    (upload_dir / saved_name).write_bytes(file_bytes)

    tag_list = [t.strip() for t in tags.split(",") if t.strip()] if tags else []

    doc = Document(
        filename=file.filename,
        file_path=saved_path,
        description=description or None,
        tags=tag_list or None,
        status="processing",
        uploaded_by=current_user.id,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    logger.info(f"文档上传成功: document_id={doc.id}, filename={file.filename}")

    # Trigger async vectorization via Celery
    try:
        from app.services.celery_tasks import vectorize_document_task
        ext = Path(file.filename).suffix.lower()
        vectorize_document_task.delay(doc.id, saved_path, ext)
        logger.info(f"已提交异步向量化任务: document_id={doc.id}")
    except Exception as e:
        logger.warning(f"Celery 任务提交失败，回退同步向量化: {e}")
        try:
            await _vectorize_document(doc, file_bytes, db)
        except Exception as ve:
            logger.error(f"文档向量化失败: doc_id={doc.id}, error={ve}")
            doc.status = "error"
            doc.error_message = str(ve)
            await db.commit()

    return DocumentUploadResponse(
        document_id=doc.id,
        status="processing",
        message="文档已上传，正在向量化处理，请稍后查询状态",
    )


@router.get("", response_model=DocumentListResponse)
async def list_documents(
    page: int = 1,
    page_size: int = 50,
    current_user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """获取文档列表（分页）。"""
    page_size = min(page_size, 100)

    count_result = await db.execute(select(func.count()).select_from(Document))
    total = count_result.scalar() or 0

    result = await db.execute(
        select(Document)
        .order_by(Document.created_at.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
    )
    docs = result.scalars().all()

    items = [
        DocumentListItem(
            id=doc.id,
            filename=doc.filename,
            description=doc.description,
            tags=doc.tags,
            status=doc.status,
            chunks_count=doc.chunks_count,
            created_at=doc.created_at,
        )
        for doc in docs
    ]

    return DocumentListResponse(items=items, total=total, page=page, page_size=page_size)


@router.get("/{document_id}", response_model=DocumentStatusResponse)
async def get_document_status(
    document_id: int,
    current_user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """查询文档处理状态。"""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise NotFoundException("文档不存在")

    return DocumentStatusResponse(
        document_id=doc.id,
        filename=doc.filename,
        description=doc.description,
        status=doc.status,
        chunks_count=doc.chunks_count,
        created_at=doc.created_at,
    )


@router.delete("/{document_id}", status_code=204)
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """删除文档。同时删除物理文件和 FAISS 向量索引。"""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise NotFoundException("文档不存在")

    await db.delete(doc)
    await db.commit()

    # Also delete the physical file
    try:
        Path(doc.file_path).unlink(missing_ok=True)
    except Exception as e:
        logger.warning(f"删除文件失败: {e}")

    # Also delete FAISS vector index
    try:
        from app.services.vector_store import delete_document as delete_faiss_index
        await delete_faiss_index(document_id)
    except Exception as e:
        logger.warning(f"删除FAISS索引失败: {e}")

    logger.info(f"文档已删除: document_id={document_id}")


@router.post("/{document_id}/reprocess", response_model=DocumentUploadResponse, status_code=202)
async def reprocess_document(
    document_id: int,
    current_user: User = Depends(get_current_admin),
    db: AsyncSession = Depends(get_db),
):
    """重新处理处于 error 状态的文档。"""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise NotFoundException("文档不存在")

    if doc.status != "error":
        raise ValidationErrorException("只有处于 error 状态的文档才能重新处理")

    doc.status = "processing"
    doc.error_message = None
    await db.commit()

    # Trigger async vectorization
    try:
        from app.services.celery_tasks import vectorize_document_task
        ext = Path(doc.file_path).suffix.lower()
        vectorize_document_task.delay(doc.id, doc.file_path, ext)
    except Exception as e:
        logger.warning(f"Celery 任务提交失败，回退同步向量化: {e}")
        file_bytes = Path(doc.file_path).read_bytes()
        try:
            await _vectorize_document(doc, file_bytes, db)
        except Exception as ve:
            logger.error(f"向量化异常: doc_id={doc.id} file={doc.file_path} — {ve}", exc_info=True)
            doc.status = "error"
            doc.error_message = str(ve)
            await db.commit()

    return DocumentUploadResponse(
        document_id=doc.id,
        status="processing",
        message="文档已重新处理",
    )


async def _vectorize_document(doc: Document, file_bytes: bytes, db: AsyncSession):
    """Extract text, chunk, embed, and store in FAISS."""
    from app.services.vector_store import add_document, chunk_text

    logger.info(f"向量化开始: doc_id={doc.id} path={doc.file_path} file_bytes_len={len(file_bytes)}")
    ext = Path(doc.file_path).suffix.lower()
    text = await _extract_doc_text(file_bytes, ext)
    logger.info(f"向量化 text提取: doc_id={doc.id} ext={ext} text_len={len(text) if text else 0}")
    if not text:
        doc.status = "error"
        doc.error_message = "无法从文档中提取文本"
        logger.error(f"向量化失败: doc_id={doc.id} ext={ext} — 无法提取文本")
        await db.commit()
        return

    # Chunk and embed
    chunks = await chunk_text(text, filename=doc.filename, ext=ext)
    if not chunks:
        doc.status = "error"
        doc.error_message = "文档内容为空或无法分块"
        logger.error(f"向量化失败: doc_id={doc.id} ext={ext} — 分块结果为空")
        await db.commit()
        return

    count = await add_document(doc.id, chunks)
    if count == 0:
        doc.status = "error"
        doc.error_message = "向量化失败"
        logger.error(f"向量化失败: doc_id={doc.id} ext={ext} — add_document 返回 0")
        await db.commit()
        return

    doc.status = "ready"
    doc.chunks_count = count
    await db.commit()
    logger.info(f"文档向量化完成: doc_id={doc.id}, chunks={count}")


async def _extract_doc_text(file_bytes: bytes, ext: str) -> str:
    """Extract text from various document formats."""
    try:
        if ext == ".pdf":
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages).strip()

        elif ext == ".docx":
            import io
            from docx import Document as DocxDocument
            doc_file = DocxDocument(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc_file.paragraphs).strip()

        elif ext in (".txt", ".md", ".py", ".js", ".ts", ".json", ".yaml", ".yml"):
            return file_bytes.decode("utf-8", errors="replace")

        else:
            # Try as plain text
            return file_bytes.decode("utf-8", errors="replace")

    except Exception as e:
        logger.error(f"文档文本提取失败: {e}")
        return ""
